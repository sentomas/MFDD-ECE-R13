import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
from scipy.signal import savgol_filter, find_peaks
import io
from docx import Document

# --- Configuration & Styling ---
st.set_page_config(
    page_title="Automotive MFDD Analyzer (Multi-Stop)",
    page_icon="🚗",
    layout="wide"
)

# --- HEADER WITH LOGO ---
col_logo, col_title = st.columns([1, 15])
with col_logo:
    st.image("https://serinthomas.co.in/favicon.ico", width=50)
with col_title:
    st.title("Passenger Vehicle MFDD Analyzer")

st.markdown("""
This tool calculates **Mean Fully Developed Deceleration (MFDD)** for passenger vehicles according to **ECE R13 / SAE J299**. 
It automatically detects multiple braking events in a single file and allows you to export the summary to Word.
""")
st.divider()

# --- Sidebar: Test Parameters ---
st.sidebar.header("1. Input Units")
time_unit = st.sidebar.selectbox("Time Unit", ["Seconds (s)", "Milliseconds (ms)"], index=0)
vel_unit = st.sidebar.selectbox("Velocity Unit", ["km/h", "m/s", "mph"], index=0)

st.sidebar.header("2. Brake Detection")
detect_max_speed = st.sidebar.checkbox("Detect Max Speed on Every Braking (Auto-Start)", value=True, help="Finds the absolute maximum speed in each cycle to perfectly set the start point of the brake.")
min_peak_speed = st.sidebar.number_input("Min Peak Speed to count as test (km/h)", value=50)
min_distance_samples = st.sidebar.number_input("Min data rows between brakes", value=100)

st.sidebar.header("3. Analysis Settings")
start_threshold_pct = st.sidebar.slider("MFDD Start Threshold", 50, 95, 80)
end_threshold_pct = st.sidebar.slider("MFDD End Threshold", 5, 40, 10)
smooth_data = st.sidebar.checkbox("Apply Smoothing Filter", value=True)
window_length = st.sidebar.slider("Smoothing Window", 5, 51, 9, step=2, disabled=not smooth_data)

# --- Helper Functions ---
def load_data(file):
    if file.name.endswith('.csv'):
        df = pd.read_csv(file)
        if len(df.columns) < 2:
            file.seek(0)
            df = pd.read_csv(file, sep=';')
        if len(df.columns) < 2:
            file.seek(0)
            df = pd.read_csv(file, sep='\t')
        return df
    else:
        return pd.read_excel(file)

def normalize_units(df, t_col, v_col, t_unit, v_unit):
    df_norm = df.copy()
    if t_unit == "Milliseconds (ms)":
        df_norm[t_col] = df_norm[t_col] / 1000.0
    if v_unit == "km/h":
        df_norm['v_ms'] = df_norm[v_col] / 3.6
        df_norm['v_kmh'] = df_norm[v_col]
    elif v_unit == "m/s":
        df_norm['v_ms'] = df_norm[v_col]
        df_norm['v_kmh'] = df_norm[v_col] * 3.6
    elif v_unit == "mph":
        df_norm['v_ms'] = df_norm[v_col] * 0.44704
        df_norm['v_kmh'] = df_norm[v_col] * 1.60934
    return df_norm

def find_nearest_idx(array, value):
    return (np.abs(array - value)).argmin()

def generate_word_report(df_results):
    doc = Document()
    doc.add_heading('MFDD Brake Test Summary Report', 0)
    
    t = doc.add_table(rows=1, cols=len(df_results.columns))
    t.style = 'Table Grid'
    
    hdr_cells = t.rows[0].cells
    for i, col_name in enumerate(df_results.columns):
        hdr_cells[i].text = str(col_name)
        
    for index, row in df_results.iterrows():
        row_cells = t.add_row().cells
        for i, col_name in enumerate(df_results.columns):
            val = row[col_name]
            if isinstance(val, float):
                val = f"{val:.2f}"
            row_cells[i].text = str(val)
            
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- Main Application ---
uploaded_file = st.file_uploader("Upload Test Data (CSV/Excel)", type=['csv', 'xlsx', 'xls'])

if uploaded_file is not None:
    try:
        raw_df = load_data(uploaded_file)
    except Exception as e:
        st.error(f"Error loading file: {e}")
        st.stop()

    cols = raw_df.columns.tolist()
    default_time_idx = 0
    default_vel_idx = 1 if len(cols) > 1 else 0

    col1, col2 = st.columns(2)
    with col1:
        t_col = st.selectbox("Select Time Column", cols, index=default_time_idx)
    with col2:
        v_col = st.selectbox("Select Velocity Column", cols, index=default_vel_idx)

    raw_df[t_col] = pd.to_numeric(raw_df[t_col], errors='coerce')
    raw_df[v_col] = pd.to_numeric(raw_df[v_col], errors='coerce')
    raw_df = raw_df.dropna(subset=[t_col, v_col]).reset_index(drop=True)

    df = normalize_units(raw_df, t_col, v_col, time_unit, vel_unit)

    if smooth_data:
        try:
            df['v_smooth_kmh'] = savgol_filter(df['v_kmh'], window_length, 3)
            df['v_smooth_ms'] = df['v_smooth_kmh'] / 3.6
        except:
            df['v_smooth_kmh'] = df['v_kmh']
            df['v_smooth_ms'] = df['v_ms']
    else:
        df['v_smooth_kmh'] = df['v_kmh']
        df['v_smooth_ms'] = df['v_ms']

    # MULTI-BRAKE DETECTION LOGIC
    peaks, _ = find_peaks(df['v_smooth_kmh'], height=min_peak_speed, distance=min_distance_samples)
    
    st.subheader(f"🏁 Detected {len(peaks)} Braking Events")

    # Calculate metrics for each peak
    results = []
    brake_dataframes = {}
    actual_start_indices = []

    for i, peak_idx in enumerate(peaks):
        next_peak_idx = peaks[i+1] if i+1 < len(peaks) else len(df)-1
        
        # --- NEW LOGIC: DETECT EXACT MAX SPEED ---
        if detect_max_speed:
            # Look slightly before the detected peak just in case the true max was shifted by noise
            search_start = max(0, peak_idx - int(min_distance_samples / 2))
            true_start_idx = df.loc[search_start:next_peak_idx, 'v_smooth_kmh'].idxmax()
        else:
            true_start_idx = peak_idx
            
        actual_start_indices.append(true_start_idx)
        v_0 = df.loc[true_start_idx, 'v_smooth_kmh']
        # -----------------------------------------

        search_window = df.loc[true_start_idx:next_peak_idx].copy().reset_index(drop=True)
        
        search_window['dt'] = search_window[t_col].diff().fillna(0)
        search_window['distance'] = (search_window['v_smooth_ms'] * search_window['dt']).cumsum()
        
        stop_indices = search_window[search_window['v_smooth_kmh'] < 1.0].index
        if len(stop_indices) > 0:
            stop_idx = stop_indices[0]
        else:
            stop_idx = search_window['v_smooth_kmh'].idxmin()

        brake_df = search_window.loc[:stop_idx].copy()
        brake_dataframes[f"Brake {i+1}"] = brake_df

        target_b_pct = start_threshold_pct / 100.0
        target_e_pct = end_threshold_pct / 100.0
        v_b_target = target_b_pct * v_0
        v_e_target = target_e_pct * v_0
        
        idx_b = find_nearest_idx(brake_df['v_smooth_kmh'].values, v_b_target)
        idx_e = find_nearest_idx(brake_df['v_smooth_kmh'].values, v_e_target)
        
        v_b = brake_df.loc[idx_b, 'v_smooth_kmh']
        v_e = brake_df.loc[idx_e, 'v_smooth_kmh']
        s_b = brake_df.loc[idx_b, 'distance']
        s_e = brake_df.loc[idx_e, 'distance']
        
        dist_interval = s_e - s_b
        mfdd = (v_b**2 - v_e**2) / (25.92 * dist_interval) if dist_interval > 0 else 0
        
        total_stop_dist = brake_df['distance'].iloc[-1]
        total_stop_time = brake_df[t_col].iloc[-1] - brake_df[t_col].iloc[0]

        results.append({
            "Test #": i + 1,
            "Initial Speed (km/h)": v_0,
            "MFDD (m/s²)": mfdd,
            "Stop Dist (m)": total_stop_dist,
            "Stop Time (s)": total_stop_time
        })

    # Plot Full Run with actual start points highlighted
    fig_raw = go.Figure()
    fig_raw.add_trace(go.Scatter(x=df[t_col], y=df['v_smooth_kmh'], mode='lines', name='Velocity (km/h)', line=dict(color='gray', width=2)))
    fig_raw.add_trace(go.Scatter(x=df.loc[actual_start_indices, t_col], y=df.loc[actual_start_indices, 'v_smooth_kmh'], mode='markers', name='Actual Starts', marker=dict(color='red', size=8, symbol='circle')))
    fig_raw.update_layout(xaxis_title="Time (s)", yaxis_title="Velocity (km/h)", hovermode="x unified", height=400)
    st.plotly_chart(fig_raw, use_container_width=True)

    # Results Table & Word Export
    if results:
        results_df = pd.DataFrame(results)
        
        st.subheader("📊 Summary Table")
        st.dataframe(results_df.style.format({"Initial Speed (km/h)": "{:.1f}", "MFDD (m/s²)": "{:.3f}", "Stop Dist (m)": "{:.2f}", "Stop Time (s)": "{:.2f}"}), use_container_width=True)

        word_file = generate_word_report(results_df)
        st.download_button(
            label="📄 Export Summary to Word",
            data=word_file,
            file_name="MFDD_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        st.divider()

        # Detailed view for a specific brake
        st.subheader("🔍 Detailed Brake View")
        selected_brake = st.selectbox("Select a brake event to view details:", list(brake_dataframes.keys()))
        
        sel_df = brake_dataframes[selected_brake]
        sel_result = results_df[results_df["Test #"] == int(selected_brake.split(" ")[1])].iloc[0]

        col_m1, col_m2, col_m3, col_m4 = st.columns(4)
        col_m1.metric("MFDD (m/s²)", f"{sel_result['MFDD (m/s²)']:.3f}")
        col_m2.metric("Initial Speed (v₀)", f"{sel_result['Initial Speed (km/h)']:.1f} km/h")
        col_m3.metric("Stopping Dist", f"{sel_result['Stop Dist (m)']:.2f} m")
        col_m4.metric("Stopping Time", f"{sel_result['Stop Time (s)']:.2f} s")

        v_0 = sel_result['Initial Speed (km/h)']
        idx_b = find_nearest_idx(sel_df['v_smooth_kmh'].values, start_threshold_pct / 100.0 * v_0)
        idx_e = find_nearest_idx(sel_df['v_smooth_kmh'].values, end_threshold_pct / 100.0 * v_0)
        t_b = sel_df.iloc[idx_b][t_col]
        t_e = sel_df.iloc[idx_e][t_col]
        v_b = sel_df.iloc[idx_b]['v_smooth_kmh']
        v_e = sel_df.iloc[idx_e]['v_smooth_kmh']

        fig = go.Figure()
        fig.add_trace(go.Scatter(x=sel_df[t_col], y=sel_df['v_smooth_kmh'], mode='lines', name='Velocity (km/h)', line=dict(color='#1f77b4', width=3)))
        fig.add_trace(go.Scatter(x=[t_b, t_e], y=[v_b, v_e], mode='markers', name='MFDD Points', marker=dict(size=12, color='red', symbol='x')))
        fig.add_shape(type="rect", x0=t_b, y0=0, x1=t_e, y1=v_0, fillcolor="red", opacity=0.1, layer="below", line_width=0)
        fig.update_layout(title=f"Zoom: {selected_brake}", xaxis_title="Time (s)", yaxis_title="Velocity (km/h)", hovermode="x unified")
        st.plotly_chart(fig, use_container_width=True)

else:
    st.info("👋 Upload a CSV/Excel file to start.")
